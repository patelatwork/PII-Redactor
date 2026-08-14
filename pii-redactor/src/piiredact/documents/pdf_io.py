"""PDF ingestion.

The corpus for this assignment arrived as a PDF export of a Word document, and
the deliverable has to be a .docx, so we need a PDF → text → .docx path.

Extraction is block-based (PyMuPDF's layout blocks) rather than line-based: a
block corresponds to a visual paragraph or table cell, which is exactly the
granularity the redactor wants as a segment.  Lines *within* a block keep their
newlines so multi-line postal addresses stay contiguous.
"""

from __future__ import annotations

import re
from pathlib import Path

#: Characters the Word → PDF converter leaves behind that break tokenisation.
_INVISIBLES = str.maketrans({
    "​": "",   # zero-width space, used as a soft column separator
    "‌": "",
    "‍": "",
    "﻿": "",
    "­": "",   # soft hyphen
    " ": " ",  # non-breaking space
    " ": " ",
    " ": " ",
})

#: An e-mail/URL wrapped across a line break: the last token of a line carries
#: an "@" or "www." but no plausible TLD, so the next line continues it.
_WRAPPED_TOKEN = re.compile(
    r"(?P<head>(?:[\w.+\-]*@[\w.\-]*|(?:https?://|www\.)[\w.\-/]*))\n[^\S\n]*(?=[\w])"
)


#: "www.example. com" -- the converter inserts a space after the dot when a
#: hostname straddles a column break.  Lowercase-only so sentence boundaries
#: ("... of the Act. In accordance ...") are left alone.
_SPLIT_TLD = re.compile(r"(?<=[\w\-])\.\s+(com|in|net|org|io|gov|edu|biz|info)\b")


def clean_text(text: str) -> str:
    text = text.translate(_INVISIBLES)
    text = _WRAPPED_TOKEN.sub(lambda m: m.group("head"), text)
    text = _SPLIT_TLD.sub(r".\1", text)
    text = re.sub(r"[^\S\n]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def merge_fragmented_blocks(blocks: list[str]) -> list[str]:
    """Re-join blocks that the extractor split mid-value.

    Dense tables -- the cover page of a prospectus is one big table -- come back
    as fragments like ``['Email:\\ncs.connect@acme.co', 'm Telephone: + 91 20',
    '45053237']``.  Left alone, each fragment is analysed separately and the
    phone number and e-mail are simply never seen.

    A block is treated as a continuation of the previous one when the previous
    block has no sentence-ending punctuation and the next starts the way a
    continuation does: lowercase, or a digit following a digit, or after a
    dangling dash.
    """
    merged: list[str] = []
    for block in blocks:
        if not merged:
            merged.append(block)
            continue
        previous = merged[-1]
        joiner = _continuation_joiner(previous, block)
        if joiner is None:
            merged.append(block)
        else:
            merged[-1] = previous + joiner + block
    return [clean_text(b) for b in merged]


def _continuation_joiner(previous: str, nxt: str) -> str | None:
    """Return the string to join two blocks with, or ``None`` to keep them apart."""
    prev_tail = previous.rstrip()
    head = nxt.lstrip()
    if not prev_tail or not head:
        return None
    if prev_tail[-1] in ".!?:;":
        return None

    last_char, first_char = prev_tail[-1], head[0]
    if prev_tail[-1] in "-–—":
        return " "
    if first_char.isdigit() and last_char.isdigit():
        return " "
    if first_char.islower() or first_char in ".,)@":
        # An e-mail or hostname broken across blocks must be rejoined with no
        # space; ordinary prose keeps its word gap.
        last_token = prev_tail.split()[-1]
        return "" if ("@" in last_token or "." in last_token) else " "
    return None


def extract_blocks(path: str | Path) -> list[str]:
    """Return one string per layout block, in reading order."""
    import pymupdf  # imported lazily; only the PDF path needs it

    blocks: list[str] = []
    with pymupdf.open(str(path)) as document:
        for page in document:
            for block in page.get_text("blocks", sort=False):
                raw = block[4]
                if not isinstance(raw, str):
                    continue
                cleaned = clean_text(raw)
                if cleaned:
                    blocks.append(cleaned)
    return merge_fragmented_blocks(blocks)


def extract_text(path: str | Path) -> str:
    return "\n\n".join(extract_blocks(path))


def pdf_to_docx(source: str | Path, destination: str | Path) -> list[str]:
    """Convert a PDF to a plain .docx and return the blocks written.

    This is a *text* conversion: it reproduces the reading-order content, not
    the original page layout, tables or images.  That is sufficient here (the
    deliverable is a redacted text document) and is stated in the README.
    """
    blocks = extract_blocks(source)
    write_blocks(blocks, destination)
    return blocks


def write_blocks(blocks: list[str], destination: str | Path) -> None:
    """Write blocks to a .docx, one paragraph each, newlines as line breaks."""
    from docx import Document

    document = Document()
    for block in blocks:
        paragraph = document.add_paragraph()
        for i, line in enumerate(block.split("\n")):
            run = paragraph.add_run(line)
            if i < block.count("\n"):
                run.add_break()
    Path(destination).parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
