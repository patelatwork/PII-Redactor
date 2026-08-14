"""Read and rewrite .docx files in place, preserving formatting.

The naive approach -- replace ``paragraph.text`` -- collapses every run and
throws away bold/italic/links/styles.  Instead we keep the run structure and
splice replacements into the runs each entity actually covers.

Runs are collected by walking the paragraph XML in document order, so runs
nested inside hyperlinks (which is how Word stores e-mail and website links,
exactly the content we most need to redact) are included.
"""

from __future__ import annotations

from collections.abc import Iterable, Iterator
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from ..redactor import Redactor
from ..types import Entity
from .segments import analyze_segments

# --------------------------------------------------------------------------
# traversal
# --------------------------------------------------------------------------


def iter_paragraphs(document: DocxDocument) -> Iterator[Paragraph]:
    """Every paragraph in the document body, tables (nested included), headers
    and footers -- i.e. every place text can hide -- each yielded exactly once.

    The de-duplication is essential, not cosmetic.  A merged table cell occupies
    several grid positions, and ``row.cells`` returns the *same* ``<w:tc>`` once
    per position: this filing yields 5,715 paragraphs from 4,639 distinct
    elements.  Without the filter the rewriter applies one paragraph's
    replacements several times over, the second pass using offsets that the
    first pass already invalidated, and the text is destroyed
    ("MEERA MOHAN IYER" -> "MEERA MOHAN IYERRA").

    Elements are held in the ``seen`` set rather than their ``id()`` so the lxml
    proxies stay alive and identity comparison stays meaningful.
    """
    seen: set = set()

    def once(paragraphs: Iterator[Paragraph]) -> Iterator[Paragraph]:
        for paragraph in paragraphs:
            element = paragraph._p
            if element in seen:
                continue
            seen.add(element)
            yield paragraph

    yield from once(_iter_block_paragraphs(document))
    for section in document.sections:
        for part in (section.header, section.footer,
                     section.even_page_header, section.even_page_footer,
                     section.first_page_header, section.first_page_footer):
            if part is not None:
                yield from once(_iter_block_paragraphs(part))


def _iter_block_paragraphs(container) -> Iterator[Paragraph]:
    yield from getattr(container, "paragraphs", [])
    for table in getattr(container, "tables", []):
        yield from _iter_table_paragraphs(table)


def _iter_table_paragraphs(table: Table) -> Iterator[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def paragraph_runs(paragraph: Paragraph) -> list[Run]:
    """Runs in document order, including those inside hyperlink wrappers."""
    return [Run(r, paragraph) for r in paragraph._p.iter(qn("w:r"))]


def paragraph_text(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph_runs(paragraph))


# --------------------------------------------------------------------------
# rewriting
# --------------------------------------------------------------------------


def rewrite_paragraph(paragraph: Paragraph, replacements: list[tuple[int, int, str]]) -> None:
    """Apply ``(start, end, new_text)`` edits to a paragraph, keeping runs.

    Edits are applied right to left so that offsets to the left of the current
    edit remain valid without recomputing the run map.
    """
    if not replacements:
        return
    runs = paragraph_runs(paragraph)
    if not runs:
        return

    spans: list[tuple[Run, int, int]] = []
    pos = 0
    for run in runs:
        spans.append((run, pos, pos + len(run.text)))
        pos += len(run.text)

    for start, end, new_text in sorted(replacements, key=lambda r: r[0], reverse=True):
        placed = False
        for run, run_start, run_end in spans:
            if run_end <= start or run_start >= end:
                continue
            local_start = max(0, start - run_start)
            local_end = min(len(run.text), end - run_start)
            head = run.text[:local_start]
            tail = run.text[local_end:] if end <= run_end else ""
            # The whole replacement goes into the first run it touches; later
            # runs only lose the characters the entity covered.
            run.text = head + ("" if placed else new_text) + tail
            placed = True


# --------------------------------------------------------------------------
# entry points
# --------------------------------------------------------------------------


def load(path: str | Path) -> DocxDocument:
    return Document(str(path))


def redact_docx(
    source: str | Path,
    destination: str | Path,
    redactor: Redactor,
) -> list[Entity]:
    """Redact ``source`` into ``destination``; returns the entities applied."""
    document = load(source)
    paragraphs = list(iter_paragraphs(document))
    segments = [paragraph_text(p) for p in paragraphs]

    entities, buckets = analyze_segments(redactor, segments)

    for paragraph, found in zip(paragraphs, buckets, strict=True):
        rewrite_paragraph(
            paragraph,
            [(e.start, e.end, redactor.surrogates.for_entity(e)) for e in found],
        )

    Path(destination).parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
    return entities


def write_docx(paragraphs: Iterable[tuple[str, str]], destination: str | Path) -> None:
    """Create a .docx from ``(style, text)`` pairs.

    Used by the PDF ingestion path, where there is no original .docx to preserve.
    """
    document = Document()
    for style, text in paragraphs:
        document.add_paragraph(text, style=style or None)
    Path(destination).parent.mkdir(parents=True, exist_ok=True)
    document.save(str(destination))
