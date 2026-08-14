"""Document IO: formatting must survive redaction, and nothing may leak."""

from __future__ import annotations

import pytest
from docx import Document
from docx.shared import Pt

from piiredact import RedactionConfig, Redactor
from piiredact.documents import redact_document
from piiredact.documents.docx_io import (
    iter_paragraphs,
    load,
    paragraph_runs,
    paragraph_text,
    rewrite_paragraph,
)
from piiredact.documents.pdf_io import clean_text, merge_fragmented_blocks

RULES_ONLY = RedactionConfig(spacy_model=None)


@pytest.fixture
def sample_docx(tmp_path):
    """A document with bold runs, a table and a header -- all the hiding places."""
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Contact Person: ")
    bold = paragraph.add_run("Rashi Prakash Patil")
    bold.bold = True
    bold.font.size = Pt(14)
    paragraph.add_run(", Managing Director.")

    document.add_paragraph("Write to rashi.patil@acme.co.in or call + 91 20 45053237.")

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Registered office"
    table.rows[0].cells[1].text = "12, Tower 2, Baner, Pune – 411 045 Maharashtra, India"

    document.sections[0].header.paragraphs[0].text = "Acme Wire Industries Limited"

    path = tmp_path / "sample.docx"
    document.save(str(path))
    return path


# --------------------------------------------------------------------------
# docx round trip
# --------------------------------------------------------------------------


def test_docx_redaction_removes_pii_from_body_table_and_header(sample_docx, tmp_path):
    out = tmp_path / "out.docx"
    redact_document(sample_docx, out, Redactor(RULES_ONLY))

    text = "\n".join(paragraph_text(p) for p in iter_paragraphs(load(out)))
    for leaked in ("Rashi", "rashi.patil", "45053237", "411 045", "Acme"):
        assert leaked not in text, f"{leaked!r} survived redaction"
    # structure preserved
    assert "Contact Person:" in text and "Managing Director" in text
    assert "Registered office" in text


def test_docx_redaction_preserves_run_formatting(sample_docx, tmp_path):
    out = tmp_path / "out.docx"
    redact_document(sample_docx, out, Redactor(RULES_ONLY))

    paragraphs = list(iter_paragraphs(load(out)))
    first = next(p for p in paragraphs if "Contact Person" in paragraph_text(p))
    bold_runs = [r for r in paragraph_runs(first) if r.bold]
    assert bold_runs, "the bold run was collapsed"
    assert "Rashi" not in bold_runs[0].text
    assert bold_runs[0].font.size == Pt(14)


def test_rewrite_paragraph_handles_entity_spanning_runs(tmp_path):
    document = Document()
    paragraph = document.add_paragraph()
    for chunk in ("Rashi ", "Prakash", " Patil rules"):
        paragraph.add_run(chunk)

    rewrite_paragraph(paragraph, [(0, 19, "John Doe")])
    assert paragraph_text(paragraph) == "John Doe rules"


def test_rewrite_paragraph_applies_multiple_edits_in_one_run():
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A then B then C")
    rewrite_paragraph(paragraph, [(0, 1, "X"), (7, 8, "Y"), (14, 15, "Z")])
    assert paragraph_text(paragraph) == "X then Y then Z"


# --------------------------------------------------------------------------
# text and pdf adapters
# --------------------------------------------------------------------------


def test_txt_in_txt_out(tmp_path):
    source = tmp_path / "in.txt"
    source.write_text(
        "Mr. Rashi Patil\n\nEmail rashi@acme.co.in\n", encoding="utf-8"
    )
    out = tmp_path / "out.txt"
    redact_document(source, out, Redactor(RULES_ONLY))
    assert "Rashi" not in out.read_text(encoding="utf-8")


def test_txt_in_docx_out(tmp_path):
    source = tmp_path / "in.txt"
    source.write_text("Mr. Rashi Patil signed.\n", encoding="utf-8")
    out = tmp_path / "out.docx"
    redact_document(source, out, Redactor(RULES_ONLY))
    text = "\n".join(paragraph_text(p) for p in iter_paragraphs(load(out)))
    assert "Rashi" not in text and "signed" in text
    assert not out.with_suffix(".source.docx").exists(), "intermediate not cleaned up"


def test_unsupported_extension_is_rejected(tmp_path):
    source = tmp_path / "in.rtf"
    source.write_text("x", encoding="utf-8")
    with pytest.raises(ValueError, match="unsupported input"):
        redact_document(source, tmp_path / "out.docx", Redactor(RULES_ONLY))


# --------------------------------------------------------------------------
# pdf extraction repair
# --------------------------------------------------------------------------


def test_clean_text_strips_zero_width_separators():
    assert clean_text("Chairman​and​Director") == "ChairmanandDirector"


def test_clean_text_rejoins_a_hostname_split_by_a_column_break():
    assert clean_text("visit www.acme. com today") == "visit www.acme.com today"


def test_merge_rejoins_a_wrapped_email_and_phone():
    blocks = ["Email: \ncs.connect@acme.co", "m Telephone: + 91 20", "45053237"]
    assert merge_fragmented_blocks(blocks) == [
        "Email: \ncs.connect@acme.com Telephone: + 91 20 45053237"
    ]


def test_merge_keeps_separate_sentences_apart():
    blocks = ["The offer closes on Thursday.", "Bidders must apply earlier."]
    assert merge_fragmented_blocks(blocks) == blocks


# --------------------------------------------------------------------------
# whole-document consistency
# --------------------------------------------------------------------------


def test_propagation_spans_the_whole_document(tmp_path):
    """A name introduced on 'page 1' is redacted where it recurs with no cue."""
    source = tmp_path / "in.txt"
    source.write_text(
        "Mr. Rashi Prakash Patil, Managing Director.\n\n"
        + "Filler paragraph about equity shares.\n\n" * 20
        + "The shares were later transferred to Rashi Prakash Patil.\n",
        encoding="utf-8",
    )
    out = tmp_path / "out.txt"
    redact_document(source, out, Redactor(RULES_ONLY))
    body = out.read_text(encoding="utf-8")
    assert "Rashi" not in body
    # the same surrogate in both places
    first = body.split("Mr. ")[1].split(",")[0]
    assert body.count(first) == 2
